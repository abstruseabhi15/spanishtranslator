import streamlit as st
import pandas as pd
from transformers import pipeline
import PyPDF2
from io import BytesIO
import docx
from tempfile import NamedTemporaryFile

# ---------------- Load Translator Model ----------------
@st.cache_resource
def get_translator():
    # Using open-source Helsinki-NLP model for Spanish->English
    return pipeline("translation_es_to_en", model="Helsinki-NLP/opus-mt-es-en")

translator = get_translator()

def translate_text(text):
    """
    Translate Spanish text to English using the loaded model.
    """
    if not text or not str(text).strip():
        return text

    text = str(text)
    max_chunk_length = 450
    chunks = []
    current_chunk = ""
    for line in text.splitlines():
        if len(current_chunk) + len(line) + 1 > max_chunk_length:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = line
        else:
            current_chunk += " " + line
    if current_chunk:
        chunks.append(current_chunk.strip())

    try:
        results = translator(chunks, max_length=512)
        translations = [r['translation_text'] for r in results]
    except Exception as e:
        translations = [f"[Translation error: {e}]"]
    return "\n\n".join(translations)

# ------------------- UI Starts Here -------------------
st.set_page_config(page_title="Spanish→English Translator", layout="wide")
st.title("🌐 Spanish → English Translator (Open Source)")
st.write("Translate Spanish text or files into English locally — no API key required.")

# ---------------- Canvas Translation ----------------
st.header("✏️ Type Spanish Text")
spanish_text = st.text_area("Enter Spanish text here:", "", height=150)

if st.button("Translate Canvas Text"):
    if spanish_text.strip():
        english_result = translate_text(spanish_text)
        st.subheader("English Translation:")
        st.text_area("Result", english_result, height=150)
        st.download_button("Download Translation (TXT)",
                           english_result.encode("utf-8"),
                           file_name="translated_canvas.txt")
        # Download as Word
        doc = docx.Document()
        for para in english_result.split('\n\n'):
            doc.add_paragraph(para)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("Download Translation (Word)", tmp.read(), file_name="translated_canvas.docx")
    else:
        st.warning("Please enter some Spanish text to translate.")

# ---------------- File Upload Translation ----------------
st.header("📂 Upload File for Translation")
uploaded_file = st.file_uploader("Upload a `.txt`, `.xlsx`, or `.pdf` file", type=['txt', 'xlsx', 'pdf'])

if uploaded_file:
    if uploaded_file.type == "text/plain":
        # Handle text files
        text = uploaded_file.read().decode('utf-8')
        st.subheader("Original (Spanish):")
        st.text_area("Preview", text, height=100)
        translated_text = translate_text(text)
        st.subheader("Translated (English):")
        st.text_area("Result", translated_text, height=100)
        # Download as TXT (same as upload)
        st.download_button("Download Translated File (TXT)",
                           translated_text.encode("utf-8"),
                           file_name="translated.txt")
        # Download as Word (universal)
        doc = docx.Document()
        for para in translated_text.split('\n\n'):
            doc.add_paragraph(para)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("Download Translated File (Word)", tmp.read(), file_name="translated.docx")

    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        # Handle Excel files
        df = pd.read_excel(uploaded_file)
        st.subheader("Original Excel (First 5 Rows):")
        st.dataframe(df.head())
        # Translate cell-by-cell
        df_translated = df.applymap(lambda x: translate_text(str(x)) if isinstance(x, str) and x.strip() else x)
        st.subheader("Translated Excel (First 5 Rows):")
        st.dataframe(df_translated.head())

        # Save translated Excel (same as upload)
        output_filename = "translated.xlsx"
        df_translated.to_excel(output_filename, index=False)
        with open(output_filename, "rb") as f:
            st.download_button("Download Translated Excel (XLSX)", f, file_name=output_filename)
        # Download as Word (universal)
        doc = docx.Document()
        for row in df_translated.itertuples(index=False):
            doc.add_paragraph("\t".join([str(cell) for cell in row]))
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("Download Translated Excel (Word)", tmp.read(), file_name="translated.xlsx.docx")

    elif uploaded_file.type == "application/pdf":
        # Handle PDF files
        pdf_bytes = uploaded_file.read()
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        st.subheader("Original PDF Text (First 1000 characters):")
        st.text_area("Preview", text[:1000], height=150)

        translated_text = translate_text(text)
        st.subheader("Translated (English):")
        st.text_area("Result", translated_text, height=150)
        # Download as TXT (since PDF creation is not supported here)
        st.download_button("Download Translated PDF (TXT)",
                           translated_text.encode("utf-8"),
                           file_name="translated_pdf.txt")
        # Download as Word (universal)
        doc = docx.Document()
        for para in translated_text.split('\n\n'):
            doc.add_paragraph(para)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button("Download Translated PDF (Word)", tmp.read(), file_name="translated_pdf.docx")
        
